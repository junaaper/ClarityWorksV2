import os
from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
import pdfplumber
try:
    import fitz  # PyMuPDF, used as a higher-fidelity PDF fallback when present.
except ImportError:
    fitz = None
from docx import Document
from PIL import Image
import pytesseract
import tempfile

import uuid

from models import ReadabilityModel
from models.simplifier import TextSimplifier
from models.rag_engine import RAGEngine
from models.concept_extractor import ConceptExtractor
from utils.change_patches import apply_changes_by_span
from utils.text_cleaner import TextCleaner

load_dotenv()

app = Flask(__name__)
CORS(app)

# Configure Tesseract path for Windows
tesseract_path = os.getenv('TESSERACT_PATH')
if tesseract_path and os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

# Initialize model
model = ReadabilityModel()
model.load_models()

# Initialize simplifier
simplifier = TextSimplifier(readability_model=model)

# Initialize RAG engine
rag_engine = RAGEngine()

# Initialize concept extractor
concept_extractor = ConceptExtractor()

# If models not found, they will use heuristic fallback
if not model.is_trained:
    print("Warning: No trained models found. Using heuristic predictions.")
    print("Run train_model.py with CLEAR Corpus to train ML models.")

@app.route('/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok',
        'model_trained': model.is_trained,
        'wordnet_available': simplifier.wordnet_available,
        'rag_answer_generation': rag_engine.llm_client is not None,
        'rag_reranker': rag_engine.ranker is not None
    })

def _pdf_text_quality_score(text):
    """Prefer PDF extraction candidates that preserve real words/glyphs."""
    if not text or not text.strip():
        return -1_000_000
    return TextCleaner.pdf_extraction_quality_metrics(text)['quality_score']


def _public_pdf_quality(metrics):
    return {
        key: value
        for key, value in metrics.items()
        if key not in {'text', 'quality_score'}
    }


def _pdf_page_record(page_number, text, extractor):
    metrics = TextCleaner.pdf_extraction_quality_metrics(text)
    cleaned_text = TextCleaner.clean_extracted_text(metrics['text'])
    return {
        'page_number': page_number,
        'text': cleaned_text,
        'extractor': extractor,
        'quality': _public_pdf_quality(metrics),
        'warnings': TextCleaner.pdf_extraction_warnings(metrics),
        '_score': metrics['quality_score'],
    }


def _best_pdf_page_record(records):
    usable = [record for record in records if record['text'] and record['text'].strip()]
    if not usable:
        return None
    return max(usable, key=lambda record: record['_score'])


def _summarize_pdf_quality(pages, page_count):
    aggregate = {
        'page_count': page_count,
        'pages_with_text': len(pages),
        'degraded_page_count': 0,
        'limited_page_count': 0,
        'replacement_char_count': 0,
        'suspicious_glyph_count': 0,
        'repaired_replacement_count': 0,
        'repaired_suspicious_glyph_count': 0,
        'private_use_char_count': 0,
        'broken_word_count': 0,
    }

    for page in pages:
        quality = page.get('quality', {})
        label = quality.get('quality_label')
        if label == 'degraded':
            aggregate['degraded_page_count'] += 1
        elif label == 'limited':
            aggregate['limited_page_count'] += 1
        for key in (
            'replacement_char_count',
            'suspicious_glyph_count',
            'repaired_replacement_count',
            'repaired_suspicious_glyph_count',
            'private_use_char_count',
            'broken_word_count',
        ):
            aggregate[key] += int(quality.get(key, 0) or 0)

    if aggregate['degraded_page_count'] > 0:
        aggregate['quality_label'] = 'degraded'
    elif aggregate['limited_page_count'] > 0 or aggregate['pages_with_text'] < page_count:
        aggregate['quality_label'] = 'limited'
    else:
        aggregate['quality_label'] = 'clean'

    warnings = []
    repaired_count = max(aggregate['repaired_replacement_count'], aggregate['repaired_suspicious_glyph_count'])
    remaining_glyph_count = max(aggregate['replacement_char_count'], aggregate['suspicious_glyph_count'])
    if repaired_count > 0:
        warnings.append(
            f"Repaired {repaired_count} likely PDF font glyph issue"
            f"{'' if repaired_count == 1 else 's'}."
        )
    if remaining_glyph_count > 0:
        warnings.append(
            f"{remaining_glyph_count} unreadable or suspicious PDF glyph"
            f"{'' if remaining_glyph_count == 1 else 's'} remain; review extracted text before using it."
        )
    if aggregate['private_use_char_count'] > 0:
        warnings.append("The PDF uses custom font glyphs that may not map cleanly to text.")
    if aggregate['pages_with_text'] < page_count:
        missing = page_count - aggregate['pages_with_text']
        warnings.append(
            f"{missing} page{'' if missing == 1 else 's'} produced no extractable text."
        )
    if aggregate['limited_page_count'] > 0:
        warnings.append(
            f"{aggregate['limited_page_count']} page{'' if aggregate['limited_page_count'] == 1 else 's'} produced limited text."
        )

    aggregate['warnings'] = warnings
    return aggregate


def _best_pdf_text_candidate(candidates):
    usable = [candidate for candidate in candidates if candidate and candidate.strip()]
    if not usable:
        return ''
    best = max(usable, key=_pdf_text_quality_score)
    return TextCleaner.pdf_extraction_quality_metrics(best)['text']


def _extract_pdfplumber_pages(pdf_path):
    """Extract PDF pages with several pdfplumber strategies and keep each best page text."""
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            page_candidates = []
            extraction_options = [
                {'expand_ligatures': True},
                {'x_tolerance': 1, 'y_tolerance': 3, 'expand_ligatures': True},
                {'x_tolerance': 2, 'y_tolerance': 3, 'expand_ligatures': True},
                {
                    'layout': True,
                    'x_tolerance': 1,
                    'y_tolerance': 3,
                    'expand_ligatures': True,
                },
            ]

            for options in extraction_options:
                try:
                    text = page.extract_text(**options) or ''
                except TypeError:
                    text = page.extract_text() or ''
                page_candidates.append(_pdf_page_record(page_number, text, 'pdfplumber'))

            try:
                deduped_page = page.dedupe_chars(tolerance=1)
                text = deduped_page.extract_text(expand_ligatures=True) or ''
                page_candidates.append(_pdf_page_record(page_number, text, 'pdfplumber'))
            except Exception:
                pass

            best = _best_pdf_page_record(page_candidates)
            if best:
                pages.append(best)

        return pages, len(pdf.pages)


def _extract_pdfplumber_text(pdf_path):
    """Backward-compatible text-only pdfplumber extraction."""
    pages, page_count = _extract_pdfplumber_pages(pdf_path)
    return '\n\n'.join(page['text'] for page in pages if page['text'].strip()), page_count


def _extract_pymupdf_pages(pdf_path):
    """Extract PDF pages with PyMuPDF when available, expanding ligatures if possible."""
    if fitz is None:
        return [], 0

    pages = []
    with fitz.open(pdf_path) as doc:
        text_flags = getattr(fitz, 'TEXTFLAGS_TEXT', 0)
        preserve_ligatures = getattr(fitz, 'TEXT_PRESERVE_LIGATURES', 0)
        expanded_ligature_flags = text_flags & ~preserve_ligatures

        for page_number, page in enumerate(doc, 1):
            page_candidates = []
            for flags in (expanded_ligature_flags, text_flags):
                try:
                    text = page.get_text('text', sort=True, flags=flags) or ''
                except TypeError:
                    text = page.get_text('text', sort=True) or ''
                page_candidates.append(_pdf_page_record(page_number, text, 'pymupdf'))

            best = _best_pdf_page_record(page_candidates)
            if best:
                pages.append(best)

        return pages, doc.page_count


def _extract_pymupdf_text(pdf_path):
    """Backward-compatible text-only PyMuPDF extraction."""
    pages, page_count = _extract_pymupdf_pages(pdf_path)
    return '\n\n'.join(page['text'] for page in pages if page['text'].strip()), page_count


def _extract_pdf_document_from_path(pdf_path):
    """Choose the highest-quality text from available PDF extractors, page by page."""
    extractor_results = []
    page_count = 0

    try:
        pages, count = _extract_pdfplumber_pages(pdf_path)
        extractor_results.append(pages)
        page_count = max(page_count, count)
    except Exception as exc:
        print(f"pdfplumber extraction warning: {exc}")

    try:
        pages, count = _extract_pymupdf_pages(pdf_path)
        extractor_results.append(pages)
        page_count = max(page_count, count)
    except Exception as exc:
        print(f"PyMuPDF extraction warning: {exc}")

    candidates_by_page = {}
    for pages in extractor_results:
        for page in pages:
            candidates_by_page.setdefault(page['page_number'], []).append(page)

    best_pages = []
    for page_number in sorted(candidates_by_page):
        best = _best_pdf_page_record(candidates_by_page[page_number])
        if best:
            page = dict(best)
            page.pop('_score', None)
            best_pages.append(page)

    full_text = '\n\n'.join(page['text'] for page in best_pages if page['text'].strip())
    quality = _summarize_pdf_quality(best_pages, page_count)
    return {
        'text': full_text,
        'page_count': page_count,
        'pages': best_pages,
        'quality': quality,
        'warnings': quality['warnings'],
    }


def _extract_pdf_text_from_path(pdf_path):
    """Backward-compatible text-only PDF extraction."""
    pdf_document = _extract_pdf_document_from_path(pdf_path)
    return pdf_document['text'], pdf_document['page_count']

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text or len(text.strip()) < 50:
            return jsonify({'error': 'Text must be at least 50 characters'}), 400

        analysis = model.predict(text)

        return jsonify({
            'success': True,
            'analysis': analysis
        })
    except Exception as e:
        print(f"Analysis error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/extract-pdf', methods=['POST'])
def extract_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        try:
            pdf_document = _extract_pdf_document_from_path(tmp_path)
        finally:
            os.unlink(tmp_path)

        # Clean the extracted text
        full_text = pdf_document['text']
        full_text = TextCleaner.remove_images_markers(full_text)
        full_text = TextCleaner.clean_extracted_text(full_text)
        full_text = TextCleaner.ensure_editable(full_text)

        return jsonify({
            'success': True,
            'text': full_text,
            'page_count': pdf_document['page_count'],
            'quality': pdf_document['quality'],
            'warnings': pdf_document['warnings']
        })
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/extract-doc', methods=['POST'])
def extract_doc():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        filename = file.filename.lower()

        if filename.endswith('.docx'):
            suffix = '.docx'
        else:
            suffix = '.doc'

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        try:
            doc = Document(tmp_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = '\n\n'.join(paragraphs)
        finally:
            os.unlink(tmp_path)

        # Clean the extracted text
        full_text = TextCleaner.remove_images_markers(full_text)
        full_text = TextCleaner.clean_extracted_text(full_text)
        full_text = TextCleaner.ensure_editable(full_text)

        return jsonify({
            'success': True,
            'text': full_text
        })
    except Exception as e:
        print(f"DOC extraction error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/extract-image', methods=['POST'])
def extract_image():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        try:
            # Open and preprocess image
            image = Image.open(tmp_path)

            # Convert to grayscale for better OCR
            if image.mode != 'L':
                image = image.convert('L')

            # Perform OCR
            text = pytesseract.image_to_string(image)

            # Get confidence data
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(c) for c in data['conf'] if int(c) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

        finally:
            os.unlink(tmp_path)

        # Clean the OCR text
        cleaned_text = TextCleaner.remove_images_markers(text)
        cleaned_text = TextCleaner.clean_extracted_text(cleaned_text)
        cleaned_text = TextCleaner.ensure_editable(cleaned_text)

        return jsonify({
            'success': True,
            'text': cleaned_text.strip(),
            'confidence': round(avg_confidence, 2)
        })
    except Exception as e:
        print(f"Image extraction error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/train', methods=['POST'])
def train_model():
    try:
        data = request.get_json()
        corpus_path = data.get('corpus_path')

        if not corpus_path:
            corpus_path = os.path.join(
                os.path.dirname(__file__),
                'data', 'clear_corpus', 'clear_corpus.csv'
            )

        if not os.path.exists(corpus_path):
            return jsonify({
                'error': f'Corpus not found at {corpus_path}'
            }), 404

        metrics = model.train(corpus_path)

        return jsonify({
            'success': True,
            'metrics': metrics
        })
    except Exception as e:
        print(f"Training error: {e}")
        return jsonify({'error': str(e)}), 500

def extract_pdf_document(file):
    """Extract page-aware text and quality metadata from a PDF file."""
    import gc, time
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name
    try:
        return _extract_pdf_document_from_path(tmp_path)
    finally:
        gc.collect()
        for _ in range(5):
            try:
                os.unlink(tmp_path)
                break
            except OSError:
                time.sleep(0.2)


def extract_text_from_pdf(file):
    """Extract text from a PDF file using the highest-quality available extractor."""
    return extract_pdf_document(file)['text']


def extract_text_from_docx(file):
    """Extract text from a DOCX file object"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n\n'.join(paragraphs)
    finally:
        os.unlink(tmp_path)
@app.route('/rag/upload', methods=['POST'])
def upload_rag_document():
    """Upload and process textbook for RAG"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        user_id = request.form.get('user_id', '')

        print(f"Uploading RAG document: {file.filename}")

        # Extract text based on file type
        filename = file.filename.lower()
        page_records = None
        extraction_quality = None
        extraction_warnings = []
        if filename.endswith('.pdf'):
            pdf_document = extract_pdf_document(file)
            extraction_quality = pdf_document['quality']
            extraction_warnings = pdf_document['warnings']
            page_records = []

            for page in pdf_document['pages']:
                page_text = TextCleaner.remove_images_markers(page['text'])
                page_text = TextCleaner.clean_textbook_text(page_text)
                page_text = TextCleaner.clean_extracted_text(page_text)
                if page_text.strip():
                    page_records.append({
                        'text': page_text,
                        'metadata': {
                            'page_number': str(page['page_number']),
                            'source_type': 'pdf',
                            'extractor': page.get('extractor', ''),
                            'extraction_quality': page.get('quality', {}).get('quality_label', ''),
                            'extraction_warnings': '; '.join(page.get('warnings', [])),
                        }
                    })

            text = '\n\n'.join(page['text'] for page in page_records)
        elif filename.endswith(('.docx', '.doc')):
            text = extract_text_from_docx(file)
        else:
            return jsonify({'error': 'Unsupported file type. Use PDF or DOCX.'}), 400

        # Clean the text before RAG processing
        text = TextCleaner.remove_images_markers(text)
        text = TextCleaner.clean_textbook_text(text)
        text = TextCleaner.clean_extracted_text(text)

        if not text or len(text) < 100:
            return jsonify({'error': 'Could not extract sufficient text from file'}), 400

        # Generate unique document ID
        doc_id = str(uuid.uuid4())

        # Upload to RAG engine
        result = rag_engine.upload_document(
            document_id=doc_id,
            text=text,
            metadata={
                'filename': file.filename,
                'user_id': user_id,
                'source_type': 'pdf' if filename.endswith('.pdf') else 'docx',
                'extraction_quality': (extraction_quality or {}).get('quality_label', ''),
                'extraction_warnings': '; '.join(extraction_warnings),
            },
            pages=page_records
        )

        return jsonify({
            'document_id': doc_id,
            'chunks_created': result['chunks_created'],
            'collection_id': result['collection_id'],
            'extraction_quality': extraction_quality,
            'warnings': extraction_warnings
        })

    except Exception as e:
        import traceback
        print(f"RAG upload error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/rag/query', methods=['POST'])
def query_rag_documents():
    """Query across uploaded textbooks with True RAG answer generation"""
    try:
        data = request.get_json()
        query_text = data['query']
        document_ids = data.get('document_ids')
        top_k = data.get('top_k', 5)

        result = rag_engine.query_documents(query_text, document_ids, top_k)

        return jsonify({
            'query': query_text,
            'answer': result['answer'],
            'sources': result['sources'],
            'has_answer': result['has_answer'],
            'results_count': len(result['sources']),
            'results': result['sources']
        })

    except Exception as e:
        print(f"RAG query error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/rag/documents/<doc_id>', methods=['DELETE'])
def delete_rag_document(doc_id):
    """Delete textbook from RAG system"""
    try:
        rag_engine.delete_document(doc_id)
        return jsonify({'status': 'deleted', 'document_id': doc_id})

    except Exception as e:
        print(f"RAG delete error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/simplify/analyze', methods=['POST'])
def analyze_for_simplification():
    """Analyze text and return suggested changes"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        target_grade_raw = data.get('target_grade', 6)
        if isinstance(target_grade_raw, str):
            import re as _re
            m = _re.search(r'\d+', target_grade_raw)
            target_grade = int(m.group()) if m else 6
        else:
            target_grade = int(target_grade_raw)
        mode = data.get('mode', 'auto')  # 'auto' or 'interactive'

        if not text or len(text.strip()) < 10:
            return jsonify({'error': 'Text must be at least 10 characters'}), 400

        result = simplifier.simplify_to_grade(text, target_grade, mode=mode)

        return jsonify({
            'original_text': text,
            'suggested_changes': result['changes'],
            'preview_text': result['simplified_text'],
            'preview_metrics': result.get('preview_metrics'),
            'target_distance': result.get('target_distance'),
            'selection_summary': result.get('selection_summary'),
        })

    except Exception as e:
        print(f"Simplification error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/simplify/apply', methods=['POST'])
def apply_selected_changes():
    """Apply only accepted changes (for interactive mode)"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        accepted_change_ids = data.get('accepted_changes', [])
        all_changes = data.get('all_changes', [])

        final_text = apply_changes_by_span(text, all_changes, accepted_change_ids)

        return jsonify({'simplified_text': final_text})

    except Exception as e:
        print(f"Apply changes error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/concepts/extract', methods=['POST'])
def extract_concepts():
    """Extract concept prerequisite graph from text or chunks"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        chunks = data.get('chunks', [])

        if chunks:
            concept_graph = concept_extractor.extract_from_chunks(chunks)
        elif text:
            concept_graph = concept_extractor.extract(text)
        else:
            return jsonify({'error': 'Provide either text or chunks'}), 400

        return jsonify({
            'success': True,
            'concept_graph': concept_graph,
        })

    except Exception as e:
        print(f"Concept extraction error: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/rag/documents/<doc_id>/chunks', methods=['GET'])
def get_document_chunks(doc_id):
    """Retrieve all text chunks for a document from ChromaDB"""
    try:
        collection_name = f"doc_{doc_id}"
        try:
            collection = rag_engine.client.get_collection(collection_name)
        except Exception:
            return jsonify({'error': 'Document not found in vector database'}), 404

        result = collection.get(include=['documents'])
        chunks = result.get('documents', [])

        return jsonify({
            'success': True,
            'chunks': chunks,
            'total_chunks': len(chunks),
        })

    except Exception as e:
        print(f"Get chunks error: {e}")
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    port = int(os.getenv('FLASK_PORT', 5001))
    debug = os.getenv('FLASK_ENV') == 'development'
    app.run(host='0.0.0.0', port=port, debug=debug)
